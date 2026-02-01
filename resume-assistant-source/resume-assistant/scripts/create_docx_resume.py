#!/usr/bin/env python3
"""
Generate DOCX resume from structured data.

Usage:
    python create_docx_resume.py output.docx --data resume_data.json

Dependencies:
    pip install python-docx
"""

import sys
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def create_resume_docx(data: dict, output_path: str) -> None:
    """
    Create a professional DOCX resume from structured data.

    Args:
        data: Resume data dictionary
        output_path: Output DOCX file path
    """
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add name (title)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", "姓名"))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(44, 62, 80)

    # Add contact info
    contact_parts = []
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("linkedin"):
        contact_parts.append(data["linkedin"])
    if data.get("github"):
        contact_parts.append(data["github"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(102, 102, 102)

    # Add summary
    if data.get("summary"):
        add_section_title(doc, "个人简介")
        summary_para = doc.add_paragraph(data["summary"])
        summary_para.paragraph_format.space_after = Pt(12)

    # Determine section order based on user status
    is_fresh_graduate = data.get("is_fresh_graduate", False)

    if is_fresh_graduate:
        # For fresh graduates: Education → Work Experience → Projects
        if data.get("education"):
            add_section_title(doc, "教育背景")
            for edu in data["education"]:
                add_education_entry(doc, edu)

        if data.get("experience"):
            add_section_title(doc, "工作经历")
            for exp in data["experience"]:
                add_experience_entry(doc, exp)

        if data.get("projects"):
            add_section_title(doc, "项目经验")
            for proj in data["projects"]:
                add_project_entry(doc, proj)
    else:
        # For experienced professionals: Work Experience → Projects → Education
        if data.get("experience"):
            add_section_title(doc, "工作经历")
            for exp in data["experience"]:
                add_experience_entry(doc, exp)

        if data.get("projects"):
            add_section_title(doc, "项目经验")
            for proj in data["projects"]:
                add_project_entry(doc, proj)

        if data.get("education"):
            add_section_title(doc, "教育背景")
            for edu in data["education"]:
                add_education_entry(doc, edu)

    # Add skills
    if data.get("skills"):
        add_section_title(doc, "技能清单")
        for skill in data["skills"]:
            skill_para = doc.add_paragraph()
            category_run = skill_para.add_run(f"{skill['category']}：")
            category_run.bold = True
            category_run.font.size = Pt(10)
            items_run = skill_para.add_run(skill["items"])
            items_run.font.size = Pt(10)

    # Add other info
    if data.get("other"):
        add_section_title(doc, "其他")
        for item in data["other"]:
            para = doc.add_paragraph(item, style="List Bullet")
            para.paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    print(f"DOCX generated: {output_path}")


def add_section_title(doc: Document, title: str) -> None:
    """Add a section title with underline."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(44, 62, 80)

    # Add border bottom (simulated with paragraph border)
    para.paragraph_format.border_bottom = True


def add_experience_entry(doc: Document, exp: dict) -> None:
    """Add a work experience entry."""
    # Header line: Company | Position | Date
    header = doc.add_paragraph()
    company_run = header.add_run(exp.get("company", ""))
    company_run.bold = True
    company_run.font.size = Pt(11)

    if exp.get("position"):
        header.add_run(" | ")
        position_run = header.add_run(exp["position"])
        position_run.font.size = Pt(11)

    if exp.get("startDate") or exp.get("endDate"):
        # Add tab and date on the right
        header.add_run("\t")
        date_text = f"{exp.get('startDate', '')} - {exp.get('endDate', '')}"
        date_run = header.add_run(date_text)
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(136, 136, 136)

    # Achievements
    if exp.get("achievements"):
        for achievement in exp["achievements"]:
            para = doc.add_paragraph(achievement, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)


def add_project_entry(doc: Document, proj: dict) -> None:
    """Add a project entry."""
    header = doc.add_paragraph()
    name_run = header.add_run(proj.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(11)

    if proj.get("role"):
        header.add_run(" | ")
        role_run = header.add_run(proj["role"])
        role_run.font.size = Pt(10)

    if proj.get("date"):
        header.add_run("\t")
        date_run = header.add_run(proj["date"])
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(136, 136, 136)

    if proj.get("tech"):
        tech_para = doc.add_paragraph()
        tech_para.add_run("技术栈：").font.size = Pt(10)
        tech_text = proj["tech"] if isinstance(proj["tech"], str) else ", ".join(proj["tech"])
        tech_run = tech_para.add_run(tech_text)
        tech_run.font.size = Pt(10)
        tech_run.font.color.rgb = RGBColor(102, 102, 102)

    if proj.get("details"):
        for detail in proj["details"]:
            para = doc.add_paragraph(detail, style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)


def add_education_entry(doc: Document, edu: dict) -> None:
    """Add an education entry."""
    header = doc.add_paragraph()
    school_run = header.add_run(edu.get("school", ""))
    school_run.bold = True
    school_run.font.size = Pt(11)

    if edu.get("degree") or edu.get("major"):
        header.add_run(" | ")
        degree_text = f"{edu.get('degree', '')} · {edu.get('major', '')}"
        degree_run = header.add_run(degree_text)
        degree_run.font.size = Pt(10)

    if edu.get("startDate") or edu.get("endDate"):
        header.add_run("\t")
        date_text = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"
        date_run = header.add_run(date_text)
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(136, 136, 136)

    if edu.get("gpa"):
        gpa_para = doc.add_paragraph()
        gpa_para.paragraph_format.left_indent = Inches(0.25)
        gpa_run = gpa_para.add_run(f"GPA: {edu['gpa']}")
        gpa_run.font.size = Pt(10)
        gpa_run.font.color.rgb = RGBColor(102, 102, 102)


def main():
    parser = argparse.ArgumentParser(description="Generate DOCX resume from JSON data")
    parser.add_argument("output", help="Output DOCX file path")
    parser.add_argument("--data", "-d", required=True, help="JSON file with resume data")

    args = parser.parse_args()

    # Load data
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"Error: Data file not found: {args.data}")
        sys.exit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    create_resume_docx(data, args.output)


if __name__ == "__main__":
    main()
