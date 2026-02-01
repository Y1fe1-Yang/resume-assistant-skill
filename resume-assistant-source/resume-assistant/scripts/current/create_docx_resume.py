#!/usr/bin/env python3
"""
Generate DOCX resume from structured data - v2.0.3 Enhanced Version

Key improvements:
- Unified font sizing across all elements
- Explicit Chinese font settings (Microsoft YaHei/SimHei)
- Consistent spacing and formatting
- Better visual hierarchy

Usage:
    python create_docx_resume.py output.docx --data resume_data.json

Dependencies:
    pip install python-docx
"""

import sys
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn  # For setting Chinese font explicitly
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ========== Font Size Constants (Unified) ==========
FONT_NAME = "Microsoft YaHei"  # Windows: 微软雅黑, can fallback to SimHei
FONT_SIZE_NAME = Pt(16)        # Name (二号字 equivalent)
FONT_SIZE_SECTION = Pt(14)     # Section titles (小二号字)
FONT_SIZE_SUBTITLE = Pt(11)    # Company/Project names (五号字)
FONT_SIZE_BODY = Pt(10.5)      # Body text (五号字)
FONT_SIZE_AUXILIARY = Pt(9)    # Contact info, dates (小五号字)

# ========== Color Constants ==========
COLOR_PRIMARY = RGBColor(44, 62, 80)      # Dark blue-gray
COLOR_SECONDARY = RGBColor(102, 102, 102) # Medium gray
COLOR_TERTIARY = RGBColor(136, 136, 136)  # Light gray


def set_chinese_font(run, font_name: str = FONT_NAME):
    """
    Explicitly set Chinese font for a Run object.
    This ensures consistent Chinese character rendering.
    """
    run.font.name = font_name
    # Set East Asian font explicitly for Chinese characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_bullet_point(paragraph):
    """Add bullet point to paragraph without using styles."""
    # Create bullet using Unicode character
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def create_resume_docx(data: dict, output_path: str) -> None:
    """
    Create a professional DOCX resume with unified font sizing.

    Args:
        data: Resume data dictionary
        output_path: Output DOCX file path
    """
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ========== Header: Name ==========
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(6)

    name_run = name_para.add_run(data.get("name", "姓名"))
    name_run.bold = True
    name_run.font.size = FONT_SIZE_NAME
    name_run.font.color.rgb = COLOR_PRIMARY
    set_chinese_font(name_run)

    # ========== Header: Contact Info ==========
    contact_parts = []
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("linkedin"):
        contact_parts.append(data["linkedin"])
    if data.get("github"):
        contact_parts.append(data["github"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(12)

        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = FONT_SIZE_AUXILIARY
        contact_run.font.color.rgb = COLOR_SECONDARY
        set_chinese_font(contact_run)

    # ========== Personal Summary ==========
    if data.get("summary"):
        add_section_title(doc, "个人简介")
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(12)
        summary_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        summary_run = summary_para.add_run(data["summary"])
        summary_run.font.size = FONT_SIZE_BODY
        set_chinese_font(summary_run)

    # ========== Determine Section Order ==========
    is_fresh_graduate = data.get("is_fresh_graduate", False)

    if is_fresh_graduate:
        # Fresh graduates: Education → Work Experience → Projects
        if data.get("education"):
            add_section_title(doc, "教育背景")
            for edu in data["education"]:
                add_education_entry(doc, edu)

        if data.get("experience"):
            add_section_title(doc, "工作经历")
            for exp in data["experience"]:
                add_experience_entry(doc, exp)

        if data.get("projects"):
            add_section_title(doc, "项目经验")
            for proj in data["projects"]:
                add_project_entry(doc, proj)
    else:
        # Experienced professionals: Work Experience → Projects → Education
        if data.get("experience"):
            add_section_title(doc, "工作经历")
            for exp in data["experience"]:
                add_experience_entry(doc, exp)

        if data.get("projects"):
            add_section_title(doc, "项目经验")
            for proj in data["projects"]:
                add_project_entry(doc, proj)

        if data.get("education"):
            add_section_title(doc, "教育背景")
            for edu in data["education"]:
                add_education_entry(doc, edu)

    # ========== Skills ==========
    if data.get("skills"):
        add_section_title(doc, "技能清单")
        for skill in data["skills"]:
            skill_para = doc.add_paragraph()
            skill_para.paragraph_format.space_after = Pt(4)
            skill_para.paragraph_format.left_indent = Inches(0.15)

            # Category (bold)
            category_run = skill_para.add_run(f"{skill['category']}：")
            category_run.bold = True
            category_run.font.size = FONT_SIZE_BODY
            set_chinese_font(category_run)

            # Items
            items_run = skill_para.add_run(skill["items"])
            items_run.font.size = FONT_SIZE_BODY
            set_chinese_font(items_run)

    # ========== Other Information ==========
    if data.get("other"):
        add_section_title(doc, "其他")
        for item in data["other"]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(4)

            # Add bullet manually using • symbol
            run = para.add_run(f"• {item}")
            run.font.size = FONT_SIZE_BODY
            set_chinese_font(run)

    doc.save(output_path)
    print(f"✅ DOCX resume generated: {output_path}")


def add_section_title(doc: Document, title: str) -> None:
    """
    Add a section title with consistent formatting and underline.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.keep_with_next = True  # Prevent orphan titles

    run = para.add_run(title)
    run.bold = True
    run.font.size = FONT_SIZE_SECTION
    run.font.color.rgb = COLOR_PRIMARY
    set_chinese_font(run)

    # Add bottom border for visual separation
    para.paragraph_format.border_bottom = True


def add_experience_entry(doc: Document, exp: dict) -> None:
    """
    Add a work experience entry with unified font sizing.
    """
    # ========== Header: Company | Position | Date ==========
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(4)

    # Company name (bold, subtitle size)
    company_run = header.add_run(exp.get("company", ""))
    company_run.bold = True
    company_run.font.size = FONT_SIZE_SUBTITLE
    set_chinese_font(company_run)

    # Position
    if exp.get("position"):
        sep_run = header.add_run(" | ")
        sep_run.font.size = FONT_SIZE_BODY
        set_chinese_font(sep_run)

        position_run = header.add_run(exp["position"])
        position_run.font.size = FONT_SIZE_BODY
        set_chinese_font(position_run)

    # Date (right-aligned using tab)
    if exp.get("startDate") or exp.get("endDate"):
        tab_run = header.add_run("\t")
        date_text = f"{exp.get('startDate', '')} - {exp.get('endDate', '')}"
        date_run = header.add_run(date_text)
        date_run.font.size = FONT_SIZE_AUXILIARY
        date_run.font.color.rgb = COLOR_TERTIARY
        set_chinese_font(date_run)

    # ========== Achievements (bullet points) ==========
    if exp.get("achievements"):
        for achievement in exp["achievements"]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)

            # Use bullet symbol
            run = para.add_run(f"• {achievement}")
            run.font.size = FONT_SIZE_BODY
            set_chinese_font(run)


def add_project_entry(doc: Document, proj: dict) -> None:
    """
    Add a project entry with unified font sizing.
    """
    # ========== Header: Project Name | Role | Date ==========
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(4)

    # Project name (bold, subtitle size)
    name_run = header.add_run(proj.get("name", ""))
    name_run.bold = True
    name_run.font.size = FONT_SIZE_SUBTITLE
    set_chinese_font(name_run)

    # Role
    if proj.get("role"):
        sep_run = header.add_run(" | ")
        sep_run.font.size = FONT_SIZE_BODY
        set_chinese_font(sep_run)

        role_run = header.add_run(proj["role"])
        role_run.font.size = FONT_SIZE_BODY
        role_run.font.color.rgb = COLOR_SECONDARY
        set_chinese_font(role_run)

    # Date
    if proj.get("date"):
        tab_run = header.add_run("\t")
        date_run = header.add_run(proj["date"])
        date_run.font.size = FONT_SIZE_AUXILIARY
        date_run.font.color.rgb = COLOR_TERTIARY
        set_chinese_font(date_run)

    # ========== Tech Stack ==========
    if proj.get("tech"):
        tech_para = doc.add_paragraph()
        tech_para.paragraph_format.space_after = Pt(4)
        tech_para.paragraph_format.left_indent = Inches(0.1)

        label_run = tech_para.add_run("技术栈：")
        label_run.font.size = FONT_SIZE_BODY
        label_run.bold = True
        set_chinese_font(label_run)

        tech_text = proj["tech"] if isinstance(proj["tech"], str) else ", ".join(proj["tech"])
        tech_run = tech_para.add_run(tech_text)
        tech_run.font.size = FONT_SIZE_BODY
        tech_run.font.color.rgb = COLOR_SECONDARY
        set_chinese_font(tech_run)

    # ========== Project Details (bullet points) ==========
    if proj.get("details"):
        for detail in proj["details"]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)

            run = para.add_run(f"• {detail}")
            run.font.size = FONT_SIZE_BODY
            set_chinese_font(run)


def add_education_entry(doc: Document, edu: dict) -> None:
    """
    Add an education entry with unified font sizing.
    """
    # ========== Header: School | Degree · Major | Date ==========
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(4)

    # School name (bold, subtitle size)
    school_run = header.add_run(edu.get("school", ""))
    school_run.bold = True
    school_run.font.size = FONT_SIZE_SUBTITLE
    set_chinese_font(school_run)

    # Degree and Major
    if edu.get("degree") or edu.get("major"):
        sep_run = header.add_run(" | ")
        sep_run.font.size = FONT_SIZE_BODY
        set_chinese_font(sep_run)

        degree_text = f"{edu.get('degree', '')} · {edu.get('major', '')}"
        degree_run = header.add_run(degree_text)
        degree_run.font.size = FONT_SIZE_BODY
        degree_run.font.color.rgb = COLOR_SECONDARY
        set_chinese_font(degree_run)

    # Date
    if edu.get("startDate") or edu.get("endDate"):
        tab_run = header.add_run("\t")
        date_text = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"
        date_run = header.add_run(date_text)
        date_run.font.size = FONT_SIZE_AUXILIARY
        date_run.font.color.rgb = COLOR_TERTIARY
        set_chinese_font(date_run)

    # ========== GPA ==========
    if edu.get("gpa"):
        gpa_para = doc.add_paragraph()
        gpa_para.paragraph_format.left_indent = Inches(0.25)
        gpa_para.paragraph_format.space_after = Pt(4)

        gpa_run = gpa_para.add_run(f"GPA: {edu['gpa']}")
        gpa_run.font.size = FONT_SIZE_BODY
        gpa_run.font.color.rgb = COLOR_SECONDARY
        set_chinese_font(gpa_run)


def main():
    parser = argparse.ArgumentParser(description="Generate DOCX resume from JSON data")
    parser.add_argument("output", help="Output DOCX file path")
    parser.add_argument("--data", "-d", required=True, help="JSON file with resume data")

    args = parser.parse_args()

    # Load data
    data_path = Path(args.data)
    if not data_path.exists():
        print(f"Error: Data file not found: {args.data}")
        sys.exit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    create_resume_docx(data, args.output)


if __name__ == "__main__":
    main()
